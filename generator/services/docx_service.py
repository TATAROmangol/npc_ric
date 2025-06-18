from docxtpl import DocxTemplate, Subdoc
import io
from grpc_clients.table_client import get_table_data
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)


def set_grid_borders(table):
    """
    Проставляет границы (сетки) для всех ячеек таблицы вручную.
    Используется для обеспечения корректного отображения таблицы.
    """
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')

            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                borders.append(border)

            tcPr.append(borders)


def generate_docx_from_template(template_bytes: bytes,
                                institution_id: int) -> bytes:
    """
    Генерирует docx-документ на основе шаблона и данных по организации.
    Вставляет таблицу, полученную через gRPC, в виде поддокумента.
    """
    logger.info(f"Generating docx for institution {institution_id}")

    # Загружаем шаблон документа из байтов
    template_stream = io.BytesIO(template_bytes)
    doc = DocxTemplate(template_stream)

    # Получаем данные для таблицы (столбцы и строки) из внешнего gRPC-сервиса
    table_data = get_table_data(institution_id)
    logger.debug(f"Table columns: {table_data['columns']}")

    if table_data["rows"]:
        logger.debug(f"First row: {table_data['rows'][0]}")
    else:
        logger.warning("No rows received from gRPC")

    # Создаём поддокумент и таблицу
    subdoc = doc.new_subdoc()
    table = subdoc.add_table(rows=1, cols=len(table_data["columns"]))

    # Заполняем заголовки таблицы
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(table_data["columns"]):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col_name)
        run.bold = True

    # Заполняем строки таблицы
    for row in table_data["rows"]:
        row_cells = table.add_row().cells
        for i in range(min(len(row), len(row_cells))):
            row_cells[i].text = row[i]
        for i in range(len(row), len(row_cells)):
            row_cells[i].text = ""

    logger.info("Applying manual borders to table cells")
    set_grid_borders(table)

    # Рендерим шаблон с подставленной таблицей
    context = {
        "table": subdoc
    }

    doc.render(context)

    # Сохраняем итоговый документ в память и возвращаем байты
    output_stream = io.BytesIO()
    doc.save(output_stream)

    logger.info("Document generated successfully")
    return output_stream.getvalue()
